"""Build docs/*.docx from docs/*.md — the readable copies, for reading away from a screen.

    python deploy/build-docx.py                 # both documents
    python deploy/build-docx.py DAILY-CHECKLIST # just one

Why this exists rather than pandoc: pandoc is not installed here and would be a
second toolchain to keep on every machine, while `python-docx` is already a
dependency away and this only has to handle the markdown these two files
actually use. That set is closed and small:

    headings, bold/italic/inline-code, links, tables, bullets, numbered lists,
    blockquotes, fenced code, horizontal rules, and `- [ ]` checkboxes

Anything else in the source is passed through as plain text rather than
silently dropped, so a new construct shows up as ugly output you can see, not
as a missing paragraph you cannot.

Two things it does on purpose:

* **Checkboxes become real ballot boxes** (☐). DAILY-CHECKLIST.md is meant to
  be ticked on paper, and `- [ ]` rendered literally is noise.
* **Tables get a header row and borders.** The playbook's tables carry the
  monetisation thresholds; a borderless run of text loses which number belongs
  to which tier.

Word keeps an exclusive lock on an open .docx and writing to one fails with
PermissionError. That is reported per file with the fix, and the other file is
still built — a half-refresh you know about beats a crash that leaves both
stale.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"

# Deep enough to be legible in Word's navigation pane, and every level these
# two documents actually use.
HEADING_SIZES = {1: 22, 2: 16, 3: 13, 4: 11.5}

ACCENT = RGBColor(0x16, 0x8A, 0x45)      # the studio's green, darkened for paper
MUTED = RGBColor(0x5A, 0x64, 0x72)
CODE_BG = RGBColor(0x1E, 0x29, 0x3B)


# ── inline markup ──────────────────────────────────────────────────────────

# One pass, one regex, alternatives ordered longest-delimiter-first so `**` is
# claimed as bold before `*` can take it as italic. Splitting on a capturing
# group keeps the delimiters in the result, which is what lets the loop below
# tell markup from text without a second scan.
INLINE = re.compile(
    r"(\*\*[^*]+\*\*"          # bold
    r"|`[^`]+`"                # inline code
    r"|\[[^\]]+\]\([^)]+\)"    # link
    r"|\*[^*\n]+\*"            # italic
    r"|_[^_\n]+_)"             # italic, underscore form
)

LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_inline(paragraph, text: str, base_size: float | None = None,
               bold_all: bool = False, bold: bool = False,
               italic: bool = False, depth: int = 0) -> None:
    """Write `text` into `paragraph`, honouring inline markdown.

    Recursive, because these documents nest markup and the flat version left
    the inner delimiters as visible characters: `**fill in the \\`00:00\\`
    line**` matched the bold alternative whole, and inserting its contents
    verbatim printed the backticks. So bold and italic recurse with the style
    carried down, and only code spans and links are terminal.

    `depth` is a cheap guard against a pathological source string; three levels
    is more nesting than any of this prose uses.
    """
    def style(run):
        if base_size:
            run.font.size = Pt(base_size)
        if bold or bold_all:
            run.bold = True
        if italic:
            run.italic = True
        return run

    for piece in INLINE.split(text):
        if not piece:
            continue
        if depth < 3 and piece.startswith("**") and piece.endswith("**"):
            add_inline(paragraph, piece[2:-2], base_size, bold_all,
                       bold=True, italic=italic, depth=depth + 1)
        elif piece.startswith("`") and piece.endswith("`"):
            run = style(paragraph.add_run(piece[1:-1]))
            run.font.name = "Consolas"
            run.font.color.rgb = CODE_BG
        elif LINK.fullmatch(piece):
            label, url = LINK.fullmatch(piece).groups()
            # The label, then the bare URL only when it adds something. A
            # footnote-per-link would double the page count of §3.
            run = style(paragraph.add_run(label))
            run.underline = True
            run.font.color.rgb = ACCENT
            if not url.startswith("#") and label not in url:
                tail = paragraph.add_run(f" ({url})")
                tail.font.size = Pt(7.5)
                tail.font.color.rgb = MUTED
        elif (depth < 3 and len(piece) > 2
                and piece[0] in "*_" and piece[-1] == piece[0]):
            add_inline(paragraph, piece[1:-1], base_size, bold_all,
                       bold=bold, italic=True, depth=depth + 1)
        else:
            style(paragraph.add_run(piece))


# ── block structure ────────────────────────────────────────────────────────

CHECKBOX = re.compile(r"^(\s*)- \[( |x|X)\]\s+(.*)$")
BULLET = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBERED = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_SEP = re.compile(r"^\s*\|?[\s:\-|]+\|[\s:\-|]*$")


RULE = re.compile(r"\s*([-*_])\1{2,}\s*")


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def starts_block(line: str) -> bool:
    """Does this line begin something new, rather than continue what precedes?"""
    return bool(not line.strip()
                or HEADING.match(line) or CHECKBOX.match(line)
                or NUMBERED.match(line) or BULLET.match(line)
                or RULE.fullmatch(line) or is_table_row(line)
                or line.strip().startswith((">", "```")))


def gather(lines: list[str], i: int, text: str) -> tuple[str, int]:
    """A list item plus its wrapped continuation lines, joined into one string.

    The markdown is hard-wrapped at 79 columns, so a single checklist item is
    often three source lines. Rendered one paragraph per line, two things break
    and both are visible on the page: inline markup spanning a wrap is left
    literal (`**Never record an IPO` with no closing `**` on that line, so no
    bold and two stray asterisks), and one tickable item becomes three
    paragraphs with only the first carrying a box.

    A continuation is any following line that does not start a block of its own.
    That is safe for these documents because they use no nested lists — a nested
    bullet matches BULLET and is therefore treated as its own item, which is
    what you would want anyway.
    """
    body = [text]
    i += 1
    while i < len(lines) and not starts_block(lines[i]):
        body.append(lines[i].strip())
        i += 1
    return " ".join(b for b in body if b), i


def cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def build(md: str, title: str) -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── fenced code: taken verbatim, including blank lines inside it
        if stripped.startswith("```"):
            i += 1
            body = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                body.append(lines[i])
                i += 1
            i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(14)
            para.paragraph_format.space_before = Pt(4)
            run = para.add_run("\n".join(body))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            continue

        # ── tables: header, separator, then rows until the block ends
        if is_table_row(line) and i + 1 < len(lines) and TABLE_SEP.match(lines[i + 1]):
            header = cells(line)
            i += 2
            body = []
            while i < len(lines) and is_table_row(lines[i]):
                body.append(cells(lines[i]))
                i += 1
            width = max(len(header), max((len(r) for r in body), default=0))
            table = doc.add_table(rows=1, cols=width)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for col, text in enumerate(header):
                cell = table.rows[0].cells[col]
                cell.text = ""
                add_inline(cell.paragraphs[0], text, base_size=9, bold_all=True)
            for row in body:
                cs = table.add_row().cells
                for col in range(width):
                    cs[col].text = ""
                    add_inline(cs[col].paragraphs[0],
                               row[col] if col < len(row) else "", base_size=9)
            doc.add_paragraph()
            continue

        # ── horizontal rule. Word has no <hr>, so a thin spacer stands in;
        #    a row of dashes would read as content.
        if RULE.fullmatch(line):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(2)
            spacer.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # ── headings
        head = HEADING.match(line)
        if head:
            level = len(head.group(1))
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(14 if level <= 2 else 9)
            para.paragraph_format.keep_with_next = True
            add_inline(para, head.group(2),
                       base_size=HEADING_SIZES.get(level, 11))
            for run in para.runs:
                run.bold = True
                if level <= 2:
                    run.font.color.rgb = ACCENT
            i += 1
            continue

        # ── blockquote: the playbook uses these for reality checks and caveats,
        #    which are the paragraphs most worth not skimming past.
        if stripped.startswith(">"):
            body = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                body.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            para.paragraph_format.right_indent = Pt(12)
            add_inline(para, " ".join(b for b in body if b), base_size=10)
            for run in para.runs:
                run.italic = True
            continue

        # ── checkbox, before the plain-bullet rule: `- [ ] x` matches both
        box = CHECKBOX.match(line)
        if box:
            indent, mark, text = box.groups()
            text, i = gather(lines, i, text)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(20 + len(indent) * 6)
            # A hanging indent so a wrapped line lines up under the text and
            # not under the box \u2014 on a printed page that is the difference
            # between a list and a wall.
            para.paragraph_format.first_line_indent = Pt(-20)
            para.paragraph_format.space_after = Pt(4)
            tick = para.add_run("\u2612  " if mark.lower() == "x" else "\u2610  ")
            tick.font.size = Pt(12)
            add_inline(para, text)
            continue

        num = NUMBERED.match(line)
        if num:
            indent, digits, text = num.groups()
            text, i = gather(lines, i, text)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(20 + len(indent) * 6)
            para.paragraph_format.first_line_indent = Pt(-20)
            para.paragraph_format.space_after = Pt(4)
            lead = para.add_run(f"{digits}.  ")
            lead.bold = True
            add_inline(para, text)
            continue

        bul = BULLET.match(line)
        if bul:
            indent, text = bul.groups()
            text, i = gather(lines, i, text)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18 + len(indent) * 6)
            para.paragraph_format.first_line_indent = Pt(-18)
            para.paragraph_format.space_after = Pt(4)
            para.add_run("\u2022  ")
            add_inline(para, text)
            continue

        # ── a paragraph: join the wrapped source lines back together, because
        #    the markdown is hard-wrapped at 79 columns and Word does its own.
        text, i = gather(lines, i, stripped)
        add_inline(doc.add_paragraph(), text)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = footer.add_run(
        f"Generated from docs/{title}.md — edit the markdown, not this file.")
    note.font.size = Pt(7.5)
    note.font.color.rgb = MUTED
    note.italic = True
    return doc


TARGETS = ["YOUTUBE-PLAYBOOK", "DAILY-CHECKLIST"]


def main(argv: list[str]) -> int:
    wanted = [a.upper().removesuffix(".MD") for a in argv] or TARGETS
    failed = 0
    for name in wanted:
        src = DOCS / f"{name}.md"
        if not src.exists():
            print(f"  {name}: no {src.relative_to(ROOT)}", file=sys.stderr)
            failed += 1
            continue
        dest = DOCS / f"{name}.docx"
        # Word's owner file, checked up front so the message names the cause
        # instead of surfacing as a bare PermissionError from python-docx.
        #
        # The name is NOT `~$` + filename. Word drops leading characters to fit
        # the old 8.3-era budget, so YOUTUBE-PLAYBOOK.docx is owned by
        # `~$UTUBE-PLAYBOOK.docx` — two characters gone, and a check for
        # `~$YOUTUBE-...` or `~$OUTUBE-...` misses it. Matching by suffix
        # instead is robust to however many it drops.
        if any(dest.name.endswith(owner.name[2:])
               for owner in DOCS.glob("~$*.docx")):
            print(f"  {name}: open in Word — close the document and re-run.",
                  file=sys.stderr)
            print(f"     (Writing now would be overwritten the moment Word "
                  f"saves.)", file=sys.stderr)
            failed += 1
            continue
        try:
            build(src.read_text(encoding="utf-8"), name).save(dest)
        except PermissionError:
            print(f"  {name}: {dest.name} is locked (open in Word?)", file=sys.stderr)
            failed += 1
            continue
        print(f"  {name}.docx  {dest.stat().st_size / 1024:.0f} KB")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
